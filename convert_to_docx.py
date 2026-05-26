#!/usr/bin/env python
"""将 论文初稿.md 转换为 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re, os

SRC = r'F:\PINN实验\venv\U-net\论文初稿.md'
DST = r'F:\PINN实验\venv\U-net\论文初稿.docx'

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_title(text, level):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(13)}
    run = p.add_run(clean_text(text))
    run.bold = True
    run.font.size = sizes.get(level, Pt(12))
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def clean_text(text):
    """去除 Markdown 标记, 保留粗体等"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold 先不去掉, 后面处理
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links
    return text

def add_body(text):
    """添加正文段落 (处理行内粗体)"""
    p = doc.add_paragraph()
    # 分割 **text** 为普通文本和粗体交替
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # 处理数学公式: $...$ 改为斜体
            sub_parts = re.split(r'(\$.+?\$)', part)
            run = p.add_run(sub_parts[0]) if sub_parts[0] else None
            for i in range(1, len(sub_parts)):
                r = p.add_run(sub_parts[i].replace('$', ''))
                r.italic = True

def add_math_paragraph(text):
    """添加单独的公式段落 (居中对齐)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    text = text.replace('$$', '').strip()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)

def parse_table(lines, start_idx):
    """解析 Markdown 表格, 返回行列表和结束索引"""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append([cell.strip() for cell in lines[i].split('|')[1:-1]])
        i += 1
    # 移除分隔行 (|---|---| 等)
    data_rows = [rows[0]] + [r for r in rows[1:] if not all(re.match(r'^-+$', c) for c in r)]
    return data_rows, i

def add_table(headers, data_rows):
    """添加 Word 表格"""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean_text(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 灰色背景
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9D9D9', qn('w:val'): 'clear'
        })
        shading.append(shd)

    # 数据行
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_text(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

# =============================================================================
# 解析 Markdown
# =============================================================================
with open(SRC, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_code_block = False
math_display = False

while i < len(lines):
    line = lines[i]

    if line.startswith('```'):
        in_code_block = not in_code_block
        i += 1
        continue

    if in_code_block:
        i += 1
        continue

    # 标题
    if line.startswith('# ') and not line.startswith('## '):
        add_title(line[2:], 1)
    elif line.startswith('## '):
        add_title(line[3:], 2)
    elif line.startswith('### '):
        add_title(line[4:], 3)

    # 公式
    elif line.strip().startswith('$$'):
        if math_display:
            math_display = False
        else:
            math_display = True
            # 取下一行
            if i + 1 < len(lines):
                add_math_paragraph(lines[i+1])
                i += 1

    # 表格
    elif line.strip().startswith('|') and not in_code_block:
        table_rows, new_i = parse_table(lines, i)
        if len(table_rows) >= 2:
            add_table(table_rows[0], table_rows[1:])
        i = new_i
        continue

    # 无序列表
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        text = re.sub(r'^[\-\*]\s+', '', line.strip())
        run = p.add_run('• ' + clean_text(text))

    # 有序列表
    elif re.match(r'^\d+\.\s', line.strip()):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        text = re.sub(r'^\d+\.\s+', '', line.strip())
        run = p.add_run(clean_text(text))

    # 空行
    elif not line.strip():
        pass

    # 正文
    else:
        # 跳过参考文献标记
        if line.strip():
            add_body(line.strip())

    i += 1

doc.save(DST)
print(f"Word 文件已保存: {DST}")
