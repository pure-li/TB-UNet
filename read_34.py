# -*- coding: utf-8 -*-
import os, glob

# Find the file
base = r'F:\PINN实验\venv\U-net'
pattern = os.path.join(base, '**', '3.4*.docx')
matches = glob.glob(pattern, recursive=True)
print('Matches:', matches)

if matches:
    path = matches[0]
    print('Reading:', repr(path))
    from docx import Document
    doc = Document(path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            print(t)
    for i, tbl in enumerate(doc.tables):
        print('--- Table', i+1, '---')
        for row in tbl.rows:
            print(' | '.join(c.text for c in row.cells))
